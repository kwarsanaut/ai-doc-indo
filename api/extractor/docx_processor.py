# api/extractors/docx_processor.py
import os
from docx import Document
from docx.shared import Inches
from typing import Dict, Any, List, Optional
from dataclasses import dataclass
import logging
import zipfile
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)

@dataclass
class DOCXExtractionResult:
    text_content: str
    metadata: Dict[str, Any]
    tables: List[List[List[str]]]
    images_found: int
    styles_info: Dict[str, Any]
    confidence_score: float
    processing_time: float

class DOCXProcessor:
    """Advanced DOCX processing with formatting and structure preservation"""
    
    def __init__(self):
        self.supported_formats = ['.docx', '.doc']
    
    def extract_content(self, file_path: str) -> DOCXExtractionResult:
        """Extract all content from DOCX file"""
        import time
        start_time = time.time()
        
        try:
            # Open document
            doc = Document(file_path)
            
            # Extract different components
            text_content = self._extract_text_content(doc)
            tables = self._extract_tables(doc)
            metadata = self._extract_metadata(doc, file_path)
            styles_info = self._extract_styles_info(doc)
            images_count = self._count_images(file_path)
            
            # Calculate confidence
            confidence = self._calculate_confidence(text_content, tables, metadata)
            
            processing_time = time.time() - start_time
            
            return DOCXExtractionResult(
                text_content=text_content,
                metadata=metadata,
                tables=tables,
                images_found=images_count,
                styles_info=styles_info,
                confidence_score=confidence,
                processing_time=processing_time
            )
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {str(e)}")
            processing_time = time.time() - start_time
            
            return DOCXExtractionResult(
                text_content="",
                metadata={"error": str(e)},
                tables=[],
                images_found=0,
                styles_info={},
                confidence_score=0.0,
                processing_time=processing_time
            )
    
    def _extract_text_content(self, doc: Document) -> str:
        """Extract text content preserving structure"""
        content_parts = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                # Check if paragraph has special formatting
                style_name = paragraph.style.name if paragraph.style else "Normal"
                
                # Add structure markers for headings
                if "Heading" in style_name:
                    level = self._extract_heading_level(style_name)
                    content_parts.append(f"\n{'#' * level} {text}\n")
                else:
                    content_parts.append(text)
        
        return '\n\n'.join(content_parts)
    
    def _extract_tables(self, doc: Document) -> List[List[List[str]]]:
        """Extract all tables from document"""
        tables = []
        
        for table in doc.tables:
            table_data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Clean cell text
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if table_data:  # Only add non-empty tables
                tables.append(table_data)
        
        return tables
    
    def _extract_metadata(self, doc: Document, file_path: str) -> Dict[str, Any]:
        """Extract document metadata and properties"""
        metadata = {}
        
        # Core properties
        core_props = doc.core_properties
        metadata.update({
            'title': getattr(core_props, 'title', ''),
            'author': getattr(core_props, 'author', ''),
            'subject': getattr(core_props, 'subject', ''),
            'category': getattr(core_props, 'category', ''),
            'comments': getattr(core_props, 'comments', ''),
            'keywords': getattr(core_props, 'keywords', ''),
            'created': str(getattr(core_props, 'created', '')),
            'modified': str(getattr(core_props, 'modified', '')),
            'last_modified_by': getattr(core_props, 'last_modified_by', ''),
            'revision': getattr(core_props, 'revision', ''),
            'version': getattr(core_props, 'version', '')
        })
        
        # Document statistics
        metadata.update({
            'paragraphs_count': len(doc.paragraphs),
            'tables_count': len(doc.tables),
            'sections_count': len(doc.sections),
            'file_size': os.path.getsize(file_path)
        })
        
        # Page setup information
        if doc.sections:
            section = doc.sections[0]
            page_setup = {
                'page_width': str(section.page_width),
                'page_height': str(section.page_height),
                'left_margin': str(section.left_margin),
                'right_margin': str(section.right_margin),
                'top_margin': str(section.top_margin),
                'bottom_margin': str(section.bottom_margin)
            }
            metadata['page_setup'] = page_setup
        
        return metadata
    
    def _extract_styles_info(self, doc: Document) -> Dict[str, Any]:
        """Extract information about document styles"""
        styles_info = {
            'styles_used': [],
            'heading_structure': [],
            'formatting_complexity': 0
        }
        
        # Track used styles
        used_styles = set()
        heading_structure = []
        
        for paragraph in doc.paragraphs:
            if paragraph.style:
                style_name = paragraph.style.name
                used_styles.add(style_name)
                
                # Track heading structure
                if "Heading" in style_name and paragraph.text.strip():
                    level = self._extract_heading_level(style_name)
                    heading_structure.append({
                        'level': level,
                        'text': paragraph.text.strip()[:100],  # First 100 chars
                        'style': style_name
                    })
        
        styles_info['styles_used'] = list(used_styles)
        styles_info['heading_structure'] = heading_structure
        styles_info['formatting_complexity'] = len(used_styles)
        
        return styles_info
    
    def _count_images(self, file_path: str) -> int:
        """Count images in DOCX file by examining the internal structure"""
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # Look for media files
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                return len(media_files)
        except Exception as e:
            logger.error(f"Failed to count images: {str(e)}")
            return 0
    
    def _extract_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name"""
        try:
            if "Heading" in style_name:
                # Extract number from "Heading 1", "Heading 2", etc.
                parts = style_name.split()
                for part in parts:
                    if part.isdigit():
                        return int(part)
            return 1
        except:
            return 1
    
    def _calculate_confidence(self, text_content: str, tables: List, metadata: Dict) -> float:
        """Calculate extraction confidence score"""
        confidence_factors = []
        
        # Text extraction success
        text_length = len(text_content.strip())
        if text_length > 0:
            confidence_factors.append(0.9)
        
        if text_length > 1000:  # Substantial content
            confidence_factors.append(0.1)
        
        # Table extraction
        if len(tables) > 0:
            confidence_factors.append(0.8)
        
        # Metadata extraction
        if metadata.get('author') or metadata.get('title'):
            confidence_factors.append(0.7)
        
        # Document structure
        if metadata.get('paragraphs_count', 0) > 0:
            confidence_factors.append(0.8)
        
        # No errors in extraction
        if 'error' not in metadata:
            confidence_factors.append(0.9)
        
        return min(1.0, sum(confidence_factors))
    
    def extract_specific_sections(self, doc: Document, section_keywords: List[str]) -> Dict[str, str]:
        """Extract specific sections based on keywords"""
        sections = {}
        current_section = None
        current_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Check if this paragraph starts a new section
            for keyword in section_keywords:
                if keyword.lower() in text.lower():
                    # Save previous section
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content)
                    
                    # Start new section
                    current_section = keyword
                    current_content = [text]
                    break
            else:
                # Continue current section
                if current_section:
                    current_content.append(text)
        
        # Save last section
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def extract_indonesian_business_info(self, doc: Document) -> Dict[str, Any]:
        """Extract Indonesian business-specific information"""
        business_info = {
            'company_names': [],
            'addresses': [],
            'phone_numbers': [],
            'email_addresses': [],
            'ktp_numbers': [],
            'npwp_numbers': []
        }
        
        import re
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Indonesian phone patterns
        phone_patterns = [
            r'\+62\s?\d{2,3}\s?\d{3,4}\s?\d{3,4}',
            r'08\d{2}\s?\d{3,4}\s?\d{3,4}',
            r'\(021\)\s?\d{3,4}\s?\d{3,4}'
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, full_text)
            business_info['phone_numbers'].extend(matches)
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        business_info['email_addresses'] = re.findall(email_pattern, full_text)
        
        # KTP pattern (16 digits)
        ktp_pattern = r'\b\d{16}\b'
        business_info['ktp_numbers'] = re.findall(ktp_pattern, full_text)
        
        # NPWP pattern
        npwp_pattern = r'\d{2}\.\d{3}\.\d{3}\.\d{1}-\d{3}\.\d{3}'
        business_info['npwp_numbers'] = re.findall(npwp_pattern, full_text)
        
        # Company indicators (PT, CV, etc.)
        company_patterns = [
            r'PT\.?\s+[A-Z][A-Za-z\s]+',
            r'CV\.?\s+[A-Z][A-Za-z\s]+',
            r'UD\.?\s+[A-Z][A-Za-z\s]+',
            r'Toko\s+[A-Z][A-Za-z\s]+'
        ]
        
        for pattern in company_patterns:
            matches = re.findall(pattern, full_text)
            business_info['company_names'].extend(matches)
        
        # Clean duplicates
        for key in business_info:
            business_info[key] = list(set(business_info[key]))
        
        return business_info
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported"""
        return any(filename.lower().endswith(fmt) for fmt in self.supported_formats)

# Example usage
if __name__ == "__main__":
    processor = DOCXProcessor()
    
    # Test with a sample DOCX
    sample_file = "sample.docx"
    if os.path.exists(sample_file):
        result = processor.extract_content(sample_file)
        print(f"Extracted {len(result.text_content)} characters")
        print(f"Found {len(result.tables)} tables")
        print(f"Found {result.images_found} images")
        print(f"Confidence: {result.confidence_score:.2f}")
        print(f"Processing time: {result.processing_time:.2f}s")
        print(f"Styles used: {len(result.styles_info['styles_used'])}")
    else:
        print("Sample DOCX file not found")
